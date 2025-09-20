# Install required packages for Google Colab
# !pip install requests beautifulsoup4 python-docx

import requests
from bs4 import BeautifulSoup
from docx import Document
import re

def scrape_coriolanus():
    """
    Scrapes Shakespeare's Coriolanus from MIT website and formats it as requested.
    """
    base_url = "http://shakespeare.mit.edu/coriolanus/"
    
    try:
        # Get the main index page
        response = requests.get(base_url + "index.html")
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Create a new document
        doc = Document()
        
        # MIT Shakespeare website structure: find all scene links
        scene_links = []
        
        # Look for links in the main content
        for link in soup.find_all('a', href=True):
            href = link['href']
            # MIT Shakespeare uses patterns like coriolanus.1.html, coriolanus.2.html, etc.
            if re.match(r'coriolanus\.\d+\.html', href):
                scene_links.append(href)
        
        # If no links found, try to construct them manually (MIT Shakespeare typically has 5 acts)
        if not scene_links:
            print("No scene links found, trying to construct URLs manually...")
            for i in range(1, 6):  # Coriolanus has 5 acts
                scene_links.append(f"coriolanus.{i}.html")
        
        # Sort the links to ensure proper order
        scene_links.sort()
        
        print(f"Processing {len(scene_links)} scenes...")
        
        for i, href in enumerate(scene_links):
            # Extract act number from href
            act_match = re.search(r'coriolanus\.(\d+)\.html', href)
            if not act_match:
                continue
            
            act_num = act_match.group(1)
            
            # Get the scene page
            scene_url = base_url + href
            print(f"Fetching: {scene_url}")
            
            try:
                scene_response = requests.get(scene_url)
                scene_response.raise_for_status()
                scene_soup = BeautifulSoup(scene_response.content, 'html.parser')
                
                # For MIT Shakespeare, each file typically contains one scene
                # But let's try to find the actual scene number
                scene_num = "1"  # Default
                
                # Look for scene information in the page
                title_elem = scene_soup.find('title')
                if title_elem:
                    title_text = title_elem.get_text()
                    scene_match = re.search(r'Scene\s+(\d+)', title_text, re.IGNORECASE)
                    if scene_match:
                        scene_num = scene_match.group(1)
                
                # Add scene heading
                heading = f"ACT {act_num} SCENE {scene_num}"
                doc.add_heading(heading, level=1)
                
                # Process the scene content
                process_scene_content(scene_soup, doc)
                
            except Exception as e:
                print(f"Error processing {scene_url}: {e}")
                continue
        
        # Save the document
        doc.save('Coriolanus_structured.docx')
        print("Successfully saved Coriolanus_structured.docx")
        
    except Exception as e:
        print(f"Error occurred: {e}")
        import traceback
        traceback.print_exc()

def process_scene_content(scene_soup, doc):
    """
    Process the content of a single scene and add it to the document.
    """
    line_number = 1
    current_speaker = None
    
    # Find the main content area
    content_area = scene_soup.find('body')
    if not content_area:
        return
    
    # Remove script and style elements
    for script in content_area(["script", "style"]):
        script.decompose()
    
    # Get all text and process it
    text_content = content_area.get_text()
    lines = text_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 3:  # Skip empty or very short lines
            continue
        
        # Skip navigation and header elements
        skip_words = ['home', 'index', 'next', 'previous', 'shakespeare', 'mit', 'the tragedy', 'coriolanus']
        if any(skip_word in line.lower() for skip_word in skip_words):
            continue
        
        # Check if this line starts with a speaker name (typically in caps)
        # Pattern: SPEAKER NAME followed by colon or period, then dialogue
        speaker_match = re.match(r'^([A-Z][A-Z\s]{1,}?)[\.:]\s*(.*)', line)
        
        if speaker_match:
            speaker = speaker_match.group(1).strip()
            dialogue = speaker_match.group(2).strip()
            
            # Clean up speaker name
            speaker = re.sub(r'[\.:]+$', '', speaker)
            speaker = re.sub(r'\s+', ' ', speaker)  # Normalize spaces
            
            # Only show speaker name when it changes
            if speaker != current_speaker:
                current_speaker = speaker
                if dialogue:
                    formatted_line = f"{line_number}: {current_speaker}: {dialogue}"
                    line_number += 1
                else:
                    # Speaker name only, no dialogue yet
                    continue
            else:
                # Same speaker, just show dialogue
                if dialogue:
                    formatted_line = f"{line_number}: {dialogue}"
                    line_number += 1
                else:
                    continue
            
            # Split long lines and add to document
            formatted_lines = split_long_line(formatted_line)
            for formatted_line_part in formatted_lines:
                doc.add_paragraph(formatted_line_part)
        
        elif current_speaker and line and not re.match(r'^[A-Z\s]+[\.:]', line):
            # Continuation of previous speaker's dialogue
            formatted_line = f"{line_number}: {line}"
            line_number += 1
            
            # Split long lines and add to document
            formatted_lines = split_long_line(formatted_line)
            for formatted_line_part in formatted_lines:
                doc.add_paragraph(formatted_line_part)

def split_long_line(line, max_length=120):
    """
    Split a line if it's longer than max_length characters.
    Returns a list of lines.
    """
    if len(line) <= max_length:
        return [line]
    
    lines = []
    current_line = ""
    words = line.split(' ')
    
    for word in words:
        if len(current_line + " " + word) <= max_length:
            if current_line:
                current_line += " " + word
            else:
                current_line = word
        else:
            if current_line:
                lines.append(current_line)
                current_line = word
            else:
                # Single word longer than max_length, just add it
                lines.append(word)
    
    if current_line:
        lines.append(current_line)
    
    return lines

# Run the scraper
if __name__ == "__main__":
    scrape_coriolanus()
